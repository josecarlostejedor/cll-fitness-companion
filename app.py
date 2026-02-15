
import streamlit as st
import pandas as pd
import base64
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

# Configuración de la página
st.set_page_config(page_title="CLL-CARE Prescripción Profesional", layout="wide", initial_sidebar_state="expanded")

# CSS Profesional para Streamlit
st.markdown("""
<style>
    .main { background-color: #f8fafc; }
    .phase-header {
        padding: 0.75rem 1.5rem;
        border-radius: 1rem;
        color: white;
        font-weight: 800;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        margin: 2rem 0 1rem 0;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    }
    .calentamiento { background: linear-gradient(90deg, #f59e0b, #d97706); }
    .resistencia { background: linear-gradient(90deg, #2563eb, #1d4ed8); }
    .enfriamiento { background: linear-gradient(90deg, #10b981, #059669); }
    
    .exercise-card {
        background: white;
        border-radius: 1.5rem;
        padding: 1.25rem;
        border: 1px solid #e2e8f0;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        margin-bottom: 1rem;
        transition: transform 0.2s;
        height: 100%;
    }
    .exercise-card:hover { transform: translateY(-3px); box-shadow: 0 10px 15px -3px rgba(0,0,0,0.1); }
    .ex-img {
        width: 100%;
        height: 140px;
        object-fit: cover;
        border-radius: 1rem;
        margin-bottom: 0.75rem;
        border: 1px solid #f1f5f9;
    }
    .ex-title { font-weight: 800; color: #1e293b; font-size: 1rem; line-height: 1.2; margin-bottom: 0.5rem; }
    .ex-badge {
        font-size: 0.7rem;
        font-weight: 700;
        background: #f1f5f9;
        color: #64748b;
        padding: 2px 8px;
        border-radius: 999px;
        display: inline-block;
        margin-bottom: 0.5rem;
    }
    .ex-info { font-size: 0.8rem; color: #475569; }
</style>
""", unsafe_allow_html=True)

def get_base64_image(image_path):
    if os.path.exists(image_path):
        with open(image_path, "rb") as img_file:
            return f"data:image/jpg;base64,{base64.b64encode(img_file.read()).decode()}"
    return "https://images.unsplash.com/photo-1571019613454-1cb2f99b2d8b?auto=format&fit=crop&q=80&w=400"

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(ns.qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

def set_cell_background(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(ns.qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

# --- BASE DE DATOS MAESTRA ---
EXERCISES_DB = {
    # SESIÓN 1 (Ya estaba correcta)
    's1_w_1': {'nombre': 'Caminar círculos + movilidad', 'img': 'images/caminar_movilidad.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's1_w_2': {'nombre': 'Equilibrio 1 pierna', 'img': 'images/equilibrio.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's1_w_3': {'nombre': 'Flexiones pared/suelo', 'img': 'images/flexiones.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '1 min'},
    's1_w_4': {'nombre': 'Sentadilla pared', 'img': 'images/sentadilla_pared.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '1 min'},
    's1_w_5': {'nombre': 'Saltar', 'img': 'images/saltar.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's1_w_6': {'nombre': 'Lanzamientos pelota', 'img': 'images/lanzamientos.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's1_r_1': {'nombre': 'Sentadilla peso corporal', 'img': 'images/sentadilla_libre.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's1_r_2': {'nombre': 'Peso muerto rumano', 'img': 'images/peso_muerto.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's1_r_3': {'nombre': 'Plancha Isométrica', 'img': 'images/plancha.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x30s'},
    's1_r_4': {'nombre': 'Press banca barra', 'img': 'images/press_banca.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's1_r_5': {'nombre': 'Curl bíceps + h', 'img': 'images/curl_hombro.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's1_r_6': {'nombre': 'Remo mancuernas', 'img': 'images/remo.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's1_c_1': {'nombre': 'Caminata + respiración', 'img': 'images/caminata_suave.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '3 min'},
    's1_c_2': {'nombre': 'Estiramiento cuádriceps', 'img': 'images/est_cuadriceps.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's1_c_3': {'nombre': 'Estiramiento isquios', 'img': 'images/est_isquios.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's1_c_4': {'nombre': 'Estiramiento pantorrilla', 'img': 'images/est_gemelos.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's1_c_5': {'nombre': 'Estiramiento bíceps', 'img': 'images/est_biceps.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's1_c_6': {'nombre': 'Estiramiento hombros', 'img': 'images/est_hombros.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's1_c_7': {'nombre': 'Movilidad cervical', 'img': 'images/movilidad_cuello.jpg', 'parte': 'Enfriamiento', 'rpe': 2, 'plan': '2 min'},
    
    # SESIÓN 2 (Restaurada)
    's2_w_1': {'nombre': '2 min caminar + movilidad', 'img': 'images/caminar_movilidad_2.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's2_w_2': {'nombre': '2 min propiocepción tobillo', 'img': 'images/propiocepcion.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's2_w_3': {'nombre': '1 min elevación rodilla+brazo', 'img': 'images/rodilla_brazo.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '1 min'},
    's2_w_4': {'nombre': '1 min sit-to-stand', 'img': 'images/sit_to_stand.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '1 min'},
    's2_w_5': {'nombre': '2 min step-up', 'img': 'images/step_up.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's2_w_6': {'nombre': '2 min boxeo suave', 'img': 'images/boxeo.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's2_r_1': {'nombre': 'Estocada adelante carga', 'img': 'images/estocada_carga.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's2_r_2': {'nombre': 'Empuje cadera', 'img': 'images/hip_thrust.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's2_r_3': {'nombre': 'Press Pallof', 'img': 'images/pallof.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's2_r_4': {'nombre': 'Press banca mancuernas', 'img': 'images/press_mancuernas.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's2_r_5': {'nombre': 'Press hombros', 'img': 'images/press_hombro_manc.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's2_r_6': {'nombre': 'Saltos sentado-parado', 'img': 'images/salto_silla.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's2_c_1': {'nombre': '3 min caminata', 'img': 'images/caminata_2.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '3 min'},
    's2_c_2': {'nombre': '1 min cuádriceps', 'img': 'images/est_cuadriceps_2.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's2_c_3': {'nombre': '1 min glúteos', 'img': 'images/est_gluteo.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's2_c_4': {'nombre': '1 min aductor', 'img': 'images/est_aductor.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's2_c_5': {'nombre': '1 min isquios', 'img': 'images/est_isquios_2.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's2_c_6': {'nombre': '1 min hombros', 'img': 'images/est_hombros_2.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's2_c_7': {'nombre': '2 min movilidad cervical', 'img': 'images/movilidad_cuello_2.jpg', 'parte': 'Enfriamiento', 'rpe': 2, 'plan': '2 min'},

    # SESIÓN 3 (Restaurada)
    's3_w_1': {'nombre': '2 min caminar + movilidad', 'img': 'images/caminar_3.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's3_w_2': {'nombre': '2 min step-ups laterales', 'img': 'images/step_lateral.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's3_w_3': {'nombre': '1 min flexiones cerradas', 'img': 'images/flexiones_cerradas.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '1 min'},
    's3_w_4': {'nombre': '1 min estocadas sitio', 'img': 'images/estocada_sitio.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '1 min'},
    's3_w_5': {'nombre': '2 min saltos tijera', 'img': 'images/jumping_jacks.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's3_w_6': {'nombre': '2 min puente glúteos', 'img': 'images/puente.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's3_r_1': {'nombre': 'Sentadilla amplia', 'img': 'images/sentadilla_sumo.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's3_r_2': {'nombre': 'Peso muerto', 'img': 'images/peso_muerto_3.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's3_r_3': {'nombre': 'Remo barra', 'img': 'images/remo_barra.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's3_r_4': {'nombre': 'Curl bíceps barra', 'img': 'images/curl_barra.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's3_r_5': {'nombre': 'Elevaciones laterales', 'img': 'images/elev_laterales.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's3_r_6': {'nombre': 'Subida caja 1 pierna', 'img': 'images/subida_caja.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's3_c_1': {'nombre': '3 min caminar', 'img': 'images/caminar_3_end.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '3 min'},
    's3_c_2': {'nombre': '1 min cuádriceps', 'img': 'images/est_cuadriceps_3.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's3_c_3': {'nombre': '1 min glúteos', 'img': 'images/est_gluteo_3.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's3_c_4': {'nombre': '1 min isquios', 'img': 'images/est_isquios_3.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's3_c_5': {'nombre': '1 min bíceps', 'img': 'images/est_biceps_3.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's3_c_6': {'nombre': '1 min hombros', 'img': 'images/est_hombros_3.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's3_c_7': {'nombre': '2 min movilidad cervical', 'img': 'images/movilidad_cuello_3.jpg', 'parte': 'Enfriamiento', 'rpe': 2, 'plan': '2 min'},

    # SESIÓN 4 (Restaurada)
    's4_w_1': {'nombre': '2 min caminar + movilidad', 'img': 'images/caminar_4.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's4_w_2': {'nombre': '2 min sit-to-stand', 'img': 'images/sit_to_stand_4.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's4_w_3': {'nombre': '1 min flexión codo banda', 'img': 'images/curl_banda.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '1 min'},
    's4_w_4': {'nombre': '1 min elevación pantorrilla', 'img': 'images/gemelos.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '1 min'},
    's4_w_5': {'nombre': '2 min saltos', 'img': 'images/saltos_cortos.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's4_w_6': {'nombre': '2 min circunducción hombros', 'img': 'images/circunduccion.jpg', 'parte': 'Calentamiento', 'rpe': 6, 'plan': '2 min'},
    's4_r_1': {'nombre': 'Estocada lateral kettlebell', 'img': 'images/estocada_lat_kb.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's4_r_2': {'nombre': 'Peso muerto', 'img': 'images/peso_muerto_4.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's4_r_3': {'nombre': 'Vuelo pecho banco inclinado', 'img': 'images/vuelo_pecho.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's4_r_4': {'nombre': 'Press tríceps supino', 'img': 'images/press_frances.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's4_r_5': {'nombre': 'Elevación frontal', 'img': 'images/elev_frontal.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's4_r_6': {'nombre': 'Sentadillas salto', 'img': 'images/sentadilla_salto.jpg', 'parte': 'Entrenamiento de Resistencia', 'rpe': 7, 'plan': '3x12'},
    's4_c_1': {'nombre': '3 min caminata', 'img': 'images/caminata_4_end.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '3 min'},
    's4_c_2': {'nombre': '1 min cuádriceps', 'img': 'images/est_cuadriceps_4.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's4_c_3': {'nombre': '1 min glúteos', 'img': 'images/est_gluteo_4.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's4_c_4': {'nombre': '1 min tríceps', 'img': 'images/est_triceps.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's4_c_5': {'nombre': '1 min pecho', 'img': 'images/est_pecho.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's4_c_6': {'nombre': '1 min hombros', 'img': 'images/est_hombros_4.jpg', 'parte': 'Enfriamiento', 'rpe': 3, 'plan': '1 min'},
    's4_c_7': {'nombre': '2 min movilidad cervical', 'img': 'images/movilidad_cuello_4.jpg', 'parte': 'Enfriamiento', 'rpe': 2, 'plan': '2 min'},
}

SESSIONS_DB = [
    {'id': 1, 'nombre': "Sesión 1: Estabilidad Base", 'ejercicios': ['s1_w_1','s1_w_2','s1_w_3','s1_w_4','s1_w_5','s1_w_6','s1_r_1','s1_r_2','s1_r_3','s1_r_4','s1_r_5','s1_r_6','s1_c_1','s1_c_2','s1_c_3','s1_c_4','s1_c_5','s1_c_6','s1_c_7']},
    {'id': 2, 'nombre': "Sesión 2: Potencia Dinámica", 'ejercicios': ['s2_w_1','s2_w_2','s2_w_3','s2_w_4','s2_w_5','s2_w_6','s2_r_1','s2_r_2','s2_r_3','s2_r_4','s2_r_5','s2_r_6','s2_c_1','s2_c_2','s2_c_3','s2_c_4','s2_c_5','s2_c_6','s2_c_7']},
    {'id': 3, 'nombre': "Sesión 3: Fuerza Integral", 'ejercicios': ['s3_w_1','s3_w_2','s3_w_3','s3_w_4','s3_w_5','s3_w_6','s3_r_1','s3_r_2','s3_r_3','s3_r_4','s3_r_5','s3_r_6','s3_c_1','s3_c_2','s3_c_3','s3_c_4','s3_c_5','s3_c_6','s3_c_7']},
    {'id': 4, 'nombre': "Sesión 4: Control y Empuje", 'ejercicios': ['s4_w_1','s4_w_2','s4_w_3','s4_w_4','s4_w_5','s4_w_6','s4_r_1','s4_r_2','s4_r_3','s4_r_4','s4_r_5','s4_r_6','s4_c_1','s4_c_2','s4_c_3','s4_c_4','s4_c_5','s4_c_6','s4_c_7']}
]

if 'rms' not in st.session_state: st.session_state.rms = {}
if 'profile' not in st.session_state: st.session_state.profile = {'nombre': '', 'apellidos': '', 'edad': 65, 'sexo': 'Hombre'}

def generate_docx(session_id):
    doc = Document()
    session = next(s for s in SESSIONS_DB if s['id'] == session_id)
    
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # Encabezado Dr. Juan Luis Sánchez...
        header = section.header
        h_para = header.paragraphs[0]
        h_para.text = "Rutina de trabajo creada por el Dr. Juan Luis Sánchez, Víctor Vicente y José Carlos Tejedor"
        h_para.style = doc.styles['Caption']
        h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Pie de página numerado
        footer = section.footer
        f_para = footer.paragraphs[0]
        f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f_run = f_para.add_run("Página ")
        add_page_number(f_run)

    doc.add_heading('PLAN TERAPÉUTICO CLL-CARE', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_patient = p_info.add_run(f"Paciente: {st.session_state.profile['nombre']} {st.session_state.profile['apellidos']} ({st.session_state.profile['sexo']})")
    run_patient.bold = True
    p_info.add_run(f"\nSesión: {session['id']} - {session['nombre']}")

    phases = [('Calentamiento', 'CALENTAMIENTO'), ('Entrenamiento de Resistencia', 'RESISTENCIA'), ('Enfriamiento', 'ENFRIAMIENTO')]
    
    for phase_key, phase_title in phases:
        ex_ids = [eid for eid in session['ejercicios'] if eid in EXERCISES_DB and EXERCISES_DB[eid]['parte'] == phase_key]
        if not ex_ids: continue
        
        doc.add_heading(phase_title, level=1)
        table = doc.add_table(rows=0, cols=3)
        table.style = 'Table Grid'
        
        for i in range(0, len(ex_ids), 3):
            row_cells = table.add_row().cells
            for j in range(3):
                if i + j < len(ex_ids):
                    eid = ex_ids[i + j]
                    ex = EXERCISES_DB[eid]
                    cell = row_cells[j]
                    
                    p_name = cell.paragraphs[0]
                    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_name.add_run(ex['nombre']).bold = True
                    
                    if os.path.exists(ex['img']):
                        p_img = cell.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # TODAS LAS IMÁGENES A 1.5 PULGADAS
                        p_img.add_run().add_picture(ex['img'], width=Inches(1.5))
                    
                    rm = st.session_state.rms.get(eid, 0.0)
                    is_autocarga = any(word in ex['nombre'].lower() for word in ['peso corporal', 'plancha', 'pared', 'equilibrio', 'sit-to-stand', 'paso', 'tijera', 'rodilla', 'boxeo', 'salto'])
                    carga = f"{int(round(rm * 0.7))} kg" if (rm > 0 and not is_autocarga) else "P.C."
                    
                    p_data = cell.add_paragraph()
                    p_data.add_run(f"Plan: {ex['plan']}\nCarga: {carga}\nRPE: {ex['rpe']}/10").font.size = Pt(8.5)

        if phase_key == 'Enfriamiento':
            doc.add_heading('INSTRUCCIONES DE REPETICIÓN', level=2)
            p_inst = doc.add_paragraph()
            p_inst.add_run("Debe repetir el protocolo completo ").font.size = Pt(11)
            p_inst.add_run("3 VECES").bold = True
            p_inst.add_run(" por sesión. Al terminar el enfriamiento, descanse 3 minutos antes de volver a empezar. Caminar 60 minutos cada día de la semana para combatir la fatiga oncológica.").font.size = Pt(11)

    # COMPROMISO + BORG
    doc.add_page_break()
    doc.add_heading('COMPROMISO DIARIO', level=1)
    img_comp = "images/compromiso_diario.jpg"
    if os.path.exists(img_comp):
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.paragraphs[-1].add_run().add_picture(img_comp, width=Inches(3.0))
    
    p_comp = doc.add_paragraph()
    p_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_c = p_comp.add_run("Caminar 60 minutos cada día de la semana para combatir la fatiga oncológica.")
    run_c.bold = True
    run_c.font.size = Pt(14)
    run_c.font.color.rgb = RGBColor(200, 0, 0)

    doc.add_paragraph("\n")
    doc.add_heading('ESCALA DE BORG', level=2)
    
    borg_data = [
        ("0", "Reposo", "D9E9FF"), ("1-2", "M. Ligero", "D9E9FF"),
        ("3-4", "Ligero", "D9FFD9"), ("5-6", "Pesado", "FFFFD9"), 
        ("7-8", "Muy Pesado", "FFE9D9"), ("9-10", "Máximo", "FFD9D9")
    ]
    borg_table = doc.add_table(rows=2, cols=6)
    borg_table.style = 'Table Grid'
    for i, (val, desc, color) in enumerate(borg_data):
        cell_top = borg_table.cell(0, i)
        set_cell_background(cell_top, color)
        p = cell_top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{val}\n").bold = True
        p.add_run(desc).font.size = Pt(7)
        borg_table.cell(1, i).height = Inches(0.4)

    target = BytesIO()
    doc.save(target)
    return target.getvalue()

# --- INTERFAZ STREAMLIT ---
st.sidebar.title("CLL-CARE ADMIN")
page = st.sidebar.radio("Navegación", ["👤 Perfil Paciente", "🏋️ Sesiones y Prescripción", "⚙️ Gestión de RM"])

if page == "👤 Perfil Paciente":
    st.title("Historial del Paciente")
    c1, c2, c3 = st.columns(3)
    with c1: st.session_state.profile['nombre'] = st.text_input("Nombre", st.session_state.profile['nombre'])
    with c2: st.session_state.profile['apellidos'] = st.text_input("Apellidos", st.session_state.profile['apellidos'])
    with c3: st.session_state.profile['sexo'] = st.selectbox("Sexo", ["Hombre", "Mujer", "Otro"])
    st.session_state.profile['edad'] = st.number_input("Edad", 1, 120, int(st.session_state.profile['edad']))
    st.info("Caminar 60 min cada día para combatir la fatiga oncológica.")

elif page == "⚙️ Gestión de RM":
    st.title("Gestión de Cargas (1RM)")
    res_exs = {k: v for k, v in EXERCISES_DB.items() if v['parte'] == 'Entrenamiento de Resistencia'}
    for eid, ex in res_exs.items():
        val = st.number_input(f"1RM {ex['nombre']} (kg)", value=float(st.session_state.rms.get(eid, 0.0)), key=f"rm_{eid}")
        st.session_state.rms[eid] = int(val)

elif page == "🏋️ Sesiones y Prescripción":
    # MODIFICACIÓN DE TÍTULOS SEGÚN SOLICITUD - SE CORRIGE 'LUIS' SIN TILDE
    st.title("Protocolos de Entrenamiento para pacientes con CLL desarrollados por el Dr. Juan Luis Sánchez")
    st.markdown("##### App creada por José Carlos Tejedor para el Proyecto BIE de Víctor García Muriel en el IES Lucía de Medrano")
    
    sid = st.radio("Seleccionar Sesión:", [s['id'] for s in SESSIONS_DB], format_func=lambda x: f"Sesión {x}", horizontal=True)
    st.download_button(label="📥 Descargar Informe Word Profesional", data=generate_docx(sid), file_name=f"Prescripcion_{st.session_state.profile['nombre']}_S{sid}.docx")
    
    active_sess = next(s for s in SESSIONS_DB if s['id'] == sid)
    st.markdown(f"### {active_sess['nombre']}")
    
    for p_key, p_label, p_class in [
        ('Calentamiento', 'CALENTAMIENTO', 'calentamiento'), 
        ('Entrenamiento de Resistencia', 'ENTRENAMIENTO DE RESISTENCIA', 'resistencia'), 
        ('Enfriamiento', 'ENFRIAMIENTO', 'enfriamiento')
    ]:
        phase_exs = [eid for eid in active_sess['ejercicios'] if eid in EXERCISES_DB and EXERCISES_DB[eid]['parte'] == p_key]
        if phase_exs:
            st.markdown(f'<div class="phase-header {p_class}">{p_label}</div>', unsafe_allow_html=True)
            cols = st.columns(3)
            for i, eid in enumerate(phase_exs):
                ex = EXERCISES_DB[eid]
                with cols[i % 3]:
                    rm = st.session_state.rms.get(eid, 0)
                    is_auto = any(w in ex['nombre'].lower() for w in ['peso corporal', 'plancha', 'pared', 'equilibrio', 'sit-to-stand'])
                    carga_text = "P.C." if is_auto else f"{int(rm*0.7)}kg"
                    
                    st.markdown(f"""
                    <div class="exercise-card">
                        <img src="{get_base64_image(ex['img'])}" class="ex-img">
                        <div class="ex-badge">{ex['parte']}</div>
                        <div class="ex-title">{ex['nombre']}</div>
                        <div class="ex-info">
                            <b>Dosis:</b> {ex['plan']}<br>
                            <b>Carga:</b> {carga_text}<br>
                            <b>RPE:</b> {ex['rpe']}/10
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
